#!/usr/bin/env python
# coding: utf-8

# In[ ]:


# Código APP CalibraData

# Librerías
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import tempfile
import os

from sklearn.isotonic import IsotonicRegression
from sklearn.metrics import roc_auc_score, brier_score_loss

from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------
# Funciones semáforo
# ---------------------------
def semaforo_auc(auc_value):
    """
    AUC ROC (discriminación)
    🟢 >= 0.80
    🟡 0.70 - 0.79
    🔴 < 0.70
    """
    if auc_value is None or (isinstance(auc_value, float) and np.isnan(auc_value)):
        return "⚪", "No disponible", "#6b7280"  # gris

    if auc_value >= 0.80:
        return "🟢", "Buena", "#16a34a"
    elif auc_value >= 0.70:
        return "🟡", "Aceptable", "#f59e0b"
    else:
        return "🔴", "Revisar", "#dc2626"


def semaforo_brier(brier_value):
    """
    Brier Score (calibración)
    🟢 < 0.10
    🟡 0.10 - 0.20
    🔴 > 0.20
    """
    if brier_value is None or (isinstance(brier_value, float) and np.isnan(brier_value)):
        return "⚪", "No disponible", "#6b7280"

    if brier_value < 0.10:
        return "🟢", "Excelente", "#16a34a"
    elif brier_value <= 0.20:
        return "🟡", "Aceptable", "#f59e0b"
    else:
        return "🔴", "Revisar", "#dc2626"


def semaforo_global(auc_value, brier_value):
    """
    Semáforo global (simple y defendible para jurado):
    🟢 Confiable: AUC >= 0.80 y Brier <= 0.20
    🔴 Revisar:   AUC < 0.70 o Brier > 0.20
    🟡 Intermedio en lo demás
    """
    # Si no hay AUC (dataset con una sola clase), no fuerzo un global “falso”
    if auc_value is None or (isinstance(auc_value, float) and np.isnan(auc_value)):
        # Igual podemos dar estado por Brier, pero lo dejo conservador
        if brier_value is not None and not (isinstance(brier_value, float) and np.isnan(brier_value)):
            if brier_value <= 0.20:
                return "🟡", "Usable con cautela (AUC no disponible)", "#f59e0b"
            else:
                return "🔴", "Revisar (AUC no disponible)", "#dc2626"
        return "⚪", "No disponible", "#6b7280"

    if (auc_value >= 0.80) and (brier_value <= 0.20):
        return "🟢", "Modelo confiable", "#16a34a"
    if (auc_value < 0.70) or (brier_value > 0.20):
        return "🔴", "Revisar modelo", "#dc2626"
    return "🟡", "Usable con cautela", "#f59e0b"


# ---------------------------
# Configuración de la app
# ---------------------------
st.set_page_config(page_title="Modelo Regresión Isotónica", layout="centered")


# ---------------------------
# CSS: ocultar icono de enlace y agrandar métricas + estilos semáforo
# ---------------------------
st.markdown("""
<style>
a.anchor-link { display: none !important; }
a[href^="#"] { display: none !important; }

.metrics-big div[data-testid="stMetricValue"] { font-size: 40px !important; }
.metrics-big div[data-testid="stMetricLabel"] { font-size: 18px !important; }

.block-container { padding-top: 2rem; }

/* Badges semáforo */
.badge {
  display: inline-flex;
  align-items: center;
  gap: .45rem;
  padding: .38rem .65rem;
  border-radius: 999px;
  font-weight: 800;
  font-size: 0.95rem;
  line-height: 1.1;
  border: 1px solid rgba(0,0,0,0.06);
  box-shadow: 0 1px 6px rgba(0,0,0,0.08);
}

.badge small{
  font-weight: 700;
  opacity: .92;
}

.badge-wrap{
  display:flex;
  gap:.6rem;
  flex-wrap: wrap;
  margin-top: .25rem;
  margin-bottom: .6rem;
}
</style>
""", unsafe_allow_html=True)


st.title("📊 CalibraData")

# Subir excel
excel_file = st.file_uploader("📂 Subir archivo Excel", type=["xlsx"])

if excel_file:

    if st.button("🧪 Procesar muestra"):

        with tempfile.TemporaryDirectory() as tmpdir:

            # Guardar Excel temporal
            excel_path = os.path.join(tmpdir, "datos.xlsx")
            with open(excel_path, "wb") as f:
                f.write(excel_file.read())

            # Datos extraídos del excel
            datos = pd.read_excel(excel_path, header=None)

            equipo = str(datos.iloc[0, 0]).strip()
            variable = str(datos.iloc[1, 0]).strip()

            X = pd.to_numeric(datos.iloc[2:, 0], errors="coerce").to_numpy()
            y = pd.to_numeric(datos.iloc[2:, 1], errors="coerce").astype(int)

            # Modelo Isotónico Pooled
            modelo = IsotonicRegression(increasing=True, out_of_bounds="clip")
            modelo.fit(X, y)

            uso_grid = np.linspace(X.min(), X.max(), 5000)
            prob_grid = modelo.predict(uso_grid)

            idx = np.where(prob_grid >= 0.5)[0]

            if len(idx) == 0:
                P50 = np.nan
                interpretacion = (
                    "Dentro del rango observado de uso acumulado, la probabilidad estimada de deriva "
                    "no alcanza el 50%, por lo que no se identifica un umbral P50 en los datos disponibles."
                )
            else:
                P50 = int(round(uso_grid[idx[0]]))
                interpretacion = (
                    f"El valor de P50 indica que, a partir de aproximadamente {P50} {variable} acumulados, "
                    "la probabilidad estimada de deriva alcanza el 50%, lo que sugiere un umbral operativo relevante "
                    "para la toma de decisiones respecto a la frecuencia de verificación o calibración."
                )

            y_prob = modelo.predict(X)

            # Métricas
            if len(np.unique(y)) > 1:
                auc = roc_auc_score(y, y_prob)
                auc_texto = f"{auc:.3f}"
                auc_num = float(auc)
            else:
                auc = None
                auc_num = None
                auc_texto = "N/A"

            brier = brier_score_loss(y, y_prob)

            # ---------------------------
            # Gráfico
            # ---------------------------
            grafico_path = os.path.join(tmpdir, "grafico.png")
            fig, ax = plt.subplots(figsize=(8, 4))

            # Datos
            ax.scatter(
                X, y,
                alpha=0.6,
                label="Observaciones",
                color="tab:blue"
            )

            # Modelo
            ax.plot(
                uso_grid, prob_grid,
                color="red",
                linewidth=2,
                label="Modelo Regresión Isotónica"
            )

            # Línea horizontal P=0.50
            ax.axhline(
                0.5,
                color="gray",
                linestyle="--",
                linewidth=1,
                label="P = 0.50"
            )

            # Línea vertical P50
            if not np.isnan(P50):
                ax.axvline(
                    P50,
                    color="black",
                    linestyle=":",
                    linewidth=2,
                    label=f"P50 ≈ {P50} {variable}"
                )

            ax.set_xlabel(f"Uso acumulado ({variable})")
            ax.set_ylabel("Probabilidad de deriva")
            ax.set_title(f"Modelo Regresión Isotónica – {equipo}")
            ax.legend()

            # ✅ Mostrar en pantalla ANTES del Word
            st.subheader("📈 Gráfico del modelo")
            st.pyplot(fig, clear_figure=False)

            # ✅ Guardar para insertarlo en el Word
            fig.savefig(grafico_path, dpi=200, bbox_inches="tight")
            plt.close(fig)

            # ---------------------------
            # Métricas del Modelo + Semáforo
            # ---------------------------
            with st.expander("📌 Ver métricas del modelo"):
                st.markdown('<div class="metrics-big">', unsafe_allow_html=True)

                c1, c2 = st.columns(2)
                c1.metric("AUC ROC", auc_texto)
                c2.metric("Brier Score", f"{brier:.3f}")

                st.markdown('</div>', unsafe_allow_html=True)

                # Semáforo individual + global
                icon_auc, label_auc, color_auc = semaforo_auc(auc_num)
                icon_br, label_br, color_br = semaforo_brier(brier)
                icon_g, label_g, color_g = semaforo_global(auc_num, brier)

                st.markdown(
                    f"""
                    <div class="badge-wrap">
                      <div class="badge" style="background:{color_auc}22; color:{color_auc};">
                        <span>{icon_auc}</span>
                        <span>AUC: <small>{label_auc}</small></span>
                      </div>
                      <div class="badge" style="background:{color_br}22; color:{color_br};">
                        <span>{icon_br}</span>
                        <span>Brier: <small>{label_br}</small></span>
                      </div>
                    </div>

                    <div class="badge" style="background:{color_g}22; color:{color_g}; width: fit-content;">
                      <span>{icon_g}</span>
                      <span><b>{label_g}</b></span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                st.markdown(
                    "*El AUC ROC refleja la capacidad del modelo para identificar correctamente casos con deriva.  \n"
                    "Mientras que el Brier Score evalúa la confiabilidad de las probabilidades estimadas.  \n"
                    "Valores altos de AUC y bajos de Brier indican un desempeño adecuado del modelo.*"
                )

            # ---------------------------
            # Creación de informe
            # ---------------------------
            plantilla = "Formulario.docx"
            doc = MailMerge(plantilla)

            doc.merge(
                equipo=equipo,
                variable=variable,
                AUC_ROC=auc_texto,
                Brier_Score=f"{brier:.3f}",
                P50="No alcanzado" if np.isnan(P50) else str(P50),
                interpretacion_tecnica=interpretacion
            )

            docx_out = os.path.join(tmpdir, "reporte.docx")
            doc.write(docx_out)

            # Insertar gráfico en Word
            d = Document(docx_out)

            for p in d.paragraphs:
                if "Gráfico" in p.text:
                    p.text = ""
                    run = p.add_run()
                    run.add_picture(grafico_path, width=Inches(4))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            d.save(docx_out)

            # Descargar Informe Word
            with open(docx_out, "rb") as f:
                st.download_button(
                    "⬇️ Descargar Informe Técnico",
                    f.read(),
                    file_name=f"Informe {equipo}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


